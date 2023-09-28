import win32com.client
import tkinter as tk
from tkinter.filedialog import askopenfilename
import docx
import re
import os

# Helper function to open a window that specifies a file's path
def cls():
    return os.system('cls')
def open_folder():
    folderpath = askopenfilename()
    return folderpath

# Helper function to check if a paragraph starts with an option (A, B, C, D)
def is_option(paragraph):
    return paragraph.startswith(("A.", "B.", "C.", "D."))

# Helper function to split options that are on the same line
def split_options(text):
    return re.split(r'\s+(?=[A-D]\.)', text)

# Initialize variables to keep track of the current question and its options

# Format an entire paragraph to readable format
def format_paragraph(doc):
    questions = []
    options = []
    document = docx.Document()
    current_options = []
    current_question = ""
    # Loop through the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Create a new paragraph with the same formatting
        new_paragraph = document.add_paragraph()
        for run in paragraph.runs:
            # Copy the runs and their formatting from the original paragraph
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.highlight_color = run.font.highlight_color
            # Check if the paragraph contains a question
            text = run.text.strip()
            if text.startswith("Câu "):
                # Save the previous question and its options if they exist
                if current_question:
                    questions.append(current_question)
                    options.append(current_options)
                # Reset the current question and options
                current_question = text
                current_options = []
            elif is_option(text):
                # Split the options if multiple are on the same line
                for option in split_options(text):
                    current_options.append(option)

    # Append the last question and its options if they exist
    if current_question:
        questions.append(current_question)
        options.append(current_options)

    # Print the extracted questions and options
    for i, question in enumerate(questions):
        document.add_paragraph(question)
        for option in options[i]:
            document.add_paragraph(option)

    document.save('temp.docx')

def extract_format_text(paragraph):
    format_text = ""
    for run in paragraph.runs:
        if run.font.highlight_color:  # Check if the text is highlighted or bold
            format_text += run.text
    return format_text

def get_correct_answer_index(options, highlights):
    for i, option_text in enumerate(options):
        for highlight in highlights:
            if option_text in highlight:
                return i + 1
    return None

def close_excel():
    file_path = os.path.abspath(r"questions.xlsx")
    try:
        print("Closing")
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False  # Optional: Hide Excel window
        workbook = excel.Workbooks.Open(file_path)
        workbook.Close(True)  # True to save changes, False to discard changes
        excel.Quit()
        cls()
    except Exception as e:
        print(f'Error: {str(e)}')