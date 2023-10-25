import pypandoc
import docx
import os, re
import win32com.client as win32
from win32com.client import constants

def doc_to_docx(file_path, del_list):
    def convert_to_docx(file_path_convert, name, del_list):
        temp_name = f'wteTemp{name}'
        document = docx.Document()
        
        # Load the DOCX document
        pypandoc.convert_file(file_path_convert, 'plain', outputfile=f'{temp_name}.txt')
        
        # Read the text from the file and replace soft returns with paragraph marks
        with open(f'{temp_name}.txt', 'r', encoding='utf-8') as file:
            text = file.readlines()
        for line in text:    
            document.add_paragraph(line.strip())
        document.save(f'{temp_name}.docx')
        #os.remove(f'{temp_name}.txt')
        del_list.append(os.path.abspath(f'{temp_name}.docx'))
        return del_list

    # Split the file path into name and extension
    name, ext = os.path.splitext(os.path.basename(file_path))
    # Rename path with .docx
    new_file_abs = f"wteTemp{name}.docx"
    if ext == ".doc":
        # Opening MS Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(file_path)
        doc.Activate ()
        # Save and Close
        word.ActiveDocument.SaveAs(new_file_abs, FileFormat=constants.wdFormatXMLDocument)
        doc.Close(False)
        del_list = convert_to_docx(file_path, name, del_list)
    elif ext == ".docx":
        del_list = convert_to_docx(file_path, name, del_list)
    else: 
        return False, del_list
    return new_file_abs, del_list


del_list = []
count = 0
new_file_path, del_list = doc_to_docx(r"C:\Users\Toan\WordToExcel\Docx\Dap an.docx", del_list)
doc = docx.Document(new_file_path)
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    text = re.sub("\x0B", "\x0A", text)
    count +=1
    print(text, count)