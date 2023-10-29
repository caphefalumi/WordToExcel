import os
import docx
import pypandoc
import pandas as pd
import win32com.client as win32
from pdf2docx import Converter
from win32com.client import constants
from utils import *

# Function to convert .doc to .docx
def format_file(file_path: str, del_list: list) -> tuple:
    """
    Format a document file to DOCX, extract formatted text, and return relevant information.

    Args:
        `file_path` (str): The path to the input document file (can be in DOC, DOCX, or PDF format).
        `del_list` (list): A list of file paths to delete after processing.

    Returns:
        Tuple[str, list, list]: A tuple containing:
            - A string representing the absolute path to the formatted DOCX file.
            - A list of extracted highlights and formatted text.
            - A list of file paths to delete after processing.

    The function performs the following tasks:
    - If the input file is in DOC format, it converts it to DOCX, extracts formatted text, and returns relevant data.
    - If the input file is in DOCX format, it extracts formatted text and returns relevant data.
    - If the input file is in PDF format, it converts it to DOCX, extracts formatted text, and returns relevant data.

    The extracted highlights are based on specific formatting rules within the document.
    """
    def convert_to_docx(file_path_convert: str, name: str, del_list: list) -> list:
        temp_name = f'wteTemp{name}'
        
        # Load the DOCX document using pypandoc
        pypandoc.convert_file(file_path_convert, 'plain', extra_args=['--wrap=none'], outputfile=f'{temp_name}.txt')
        document = docx.Document()
        
        # Read the text from the file and replace soft returns with paragraph marks
        with open(f'{temp_name}.txt', 'r', encoding='utf-8') as file:
            text = file.readlines()
        
        for line in text:
            document.add_paragraph(line)
        
        document.save(f'{temp_name}.docx')
        os.remove(f'{temp_name}.txt')
        del_list.append(os.path.abspath(f'{temp_name}.docx'))
        return del_list

    # Function to extract formatted text
    def extract_original_format(file_path: str) -> list:
        highlights = []
        document = docx.Document(file_path)
        #Append the highlighted text
        for paragraph in document.paragraphs:
            highlighted_text = extract_format_text(paragraph)
            if is_option(highlighted_text): 
                highlights.append(CFL(re.sub(r'^[a-dA-D]\.', '', highlighted_text).strip()))
        return highlights

    # Split the file path into name and extension
    name, ext = os.path.splitext(os.path.basename(file_path))
    abs_file_path = os.path.abspath(file_path)
    new_abs_file_path = os.path.splitext(abs_file_path)[0] + '.docx'
    new_file_abs_name = f"wteTemp{name}"
    # If the extension is .doc change it into .docx and do the same as .docx
    if ext == ".doc":
        temp_name = f"wteDocTemp{name}"
        temp_path = os.path.abspath(f"{temp_name}.docx")
        
        # Opening MS Word
        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(abs_file_path)
        doc.Activate()
        
        # Save and Close
        word.ActiveDocument.SaveAs(temp_name, FileFormat=constants.wdFormatXMLDocument)
        doc.Close(False)
        
        # Get the list to delete
        del_list = convert_to_docx(temp_path, name, del_list)
        del_list.append(temp_path)
        highlights = extract_original_format(temp_path)
        return temp_path, highlights, del_list

    elif ext == ".docx":
        del_list = convert_to_docx(, name, del_list)
        highlights = extract_original_format(file_path)
        return new_abs_file_path, highlights, del_list   
    elif ext == ".pdf":
        cv = Converter(abs_file_path)
        cv.convert(new_abs_file_path, start=0, end=None)
        cv.close()
        del_list = convert_to_docx(new_abs_file_path, name, del_list)
        highlights = extract_original_format(new_abs_file_path)
        return new_abs_file_path, highlights, del_list
    return False, None, None

# Function to process questions and options
def question_create(doc, current_question: str, current_options: list, highlights: list, data: list, platform: str, selected_options: list, question_numbers: int) -> int:
    """
    Process a document to create quiz questions and options based on specific formatting.

    Args:
        `doc`: The document to process, typically a collection of paragraphs.
        `current_question` (str): The current question text.
        `current_options` (list): A list of current answer options.
        `highlights` (list): A list of highlighted text indicating correct answers.
        `data` (list): The data list to which quiz questions and options will be added.
        `platform` (str): The selected platform for the quiz.
        `selected_options` (list): The selected options for quiz formatting.
        `question_numbers` (int): The current question number being processed.

    Returns:
        int: The updated question number after processing.

    This function processes the paragraphs within a document, identifies questions and options based on specific formatting rules, and creates quiz questions with answer options and highlights. It updates the data list with the generated quiz questions.

    When processing a document:
    - A new question is identified when the text formatting indicates a question.
    - Options are added to the current question as they are encountered.
    - Once a new question is detected, the previous question is completed, and the data list is updated.

    After processing, the function also handles the last question and returns the updated question number.

    Note: The document structure and formatting rules must align with the processing logic for accurate results.
    """
    def last_question(current_question: str, current_options: list, highlights: list, data: list, platform: str, selected_options: list, question_numbers: int) -> int:
        if current_question and len(current_options) > 0:
            question_numbers += 1
            current_question, current_options = process_options(current_question, current_options, selected_options, question_numbers)
            create_quiz(data, current_question, current_options, highlights, platform)
        return question_numbers
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if is_question(text):
            if current_question and len(current_options) > 0:
                current_question, current_options = process_options(current_question, current_options, selected_options, question_numbers)
                question_numbers += 1
                create_quiz(data, current_question, current_options, highlights, platform)
            current_question = text
            current_options.clear()  # Clear the options list for the new questions
        elif is_option(text):
            for option in split_options(text):
                current_options.append(option)

    question_numbers = last_question(current_question, current_options, highlights, data, platform, selected_options, question_numbers)
    # Function to process the last question

    return question_numbers

# Function to create an Excel data frame
def data_frame(data: list, file_path: str, selected_options: list, open: bool = True) -> None:
    """
    Convert a list of data into a DataFrame, optionally shuffle rows, and save it as an Excel file.

    Args:
        `data` (list): The data to be converted into a DataFrame.
        `file_path` (str): The path to the input file for naming the output Excel file.
        `selected_options` (list): A list of options that may include "Xáo trộn câu hỏi" to shuffle rows.
    """
    output_directory = "Output"
    
    # Create the output directory if it doesn't exist
    os.makedirs(output_directory, exist_ok=True)
    
    # Get the file name without extension
    file_name = os.path.splitext(os.path.basename(file_path))[0] + ".xlsx"
    output_path = os.path.join(output_directory, file_name)
    
    df = pd.DataFrame(data)
    
    if "Xáo trộn câu hỏi" in selected_options:
        df = df.sample(frac=1)  # frac=1 shuffles all rows randomly
    # Open the output directory
    df.to_excel(output_path, index=False)
    if open:
        os.startfile(output_path)
